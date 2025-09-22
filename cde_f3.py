from docx import Document
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline
from docx.shared import Pt
from docx.oxml.ns import qn
import json
import os
from sentence_transformers import SentenceTransformer
import faiss
from rank_bm25 import BM25Okapi
from collections import defaultdict
import numpy as np
import re
import torch

# 벡터검색(텍스트 청크 임베딩 & FAISS를 이용 )
class TextIndexer:
    def __init__(self, model_name="jxm/cde-small-v2"):
        self.model = SentenceTransformer(model_name)
        self.idx = None
        self.chunks = []

    def chunk(self, segments, max_chars=1000):
        res = []
        for seg in segments:
            buf, acc = [], 0
            for line in seg["text"].split("\n"):
                if not line.strip():
                    continue
                if acc + len(line) > max_chars and buf:
                    res.append({"title": seg["title"], "text": " ".join(buf)})
                    buf, acc = [], 0
                buf.append(line.strip()); acc += len(line)
            if buf:
                res.append({"title": seg["title"], "text": " ".join(buf)})
        self.chunks = res
        return res

    # chunks에 대한 벡터 인덱스 구축 (cde_minicorpus.pt 활용)
    def build(self, embeddings_file="cde_minicorpus.pt"):
        try:
            embeddings_data = torch.load(embeddings_file)
            embs = embeddings_data['embeddings'].numpy().astype("float32")
            self.chunks = embeddings_data['chunks']

            dim = embs.shape[1]
            self.idx = faiss.IndexFlatIP(dim)
            self.idx.add(embs)
            print("Document embeddings loaded and FAISS index built successfully from .pt file.")
        except FileNotFoundError:
            print(f"Error: {embeddings_file} not found. Building embeddings from scratch.")
            texts = [c["text"] for c in self.chunks]
            if not texts:
                self.idx = None
                return
            embs = self.model.encode(texts, normalize_embeddings=True)
            dim = embs.shape[1]
            self.idx = faiss.IndexFlatIP(dim)
            self.idx.add(embs.astype("float32"))

    # 주어진 쿼리에 대해 벡터 유사성 검색 수행 (cde_query_emb.pt 활용)
    def search_dense(self, query: str, topk=4, query_file="cde_query_emb.pt"):
        if self.idx is None or not self.chunks:
            return []
        
        try:
            # 미리 계산된 쿼리 임베딩 파일 로드
            query_embs = torch.load(query_file)
            qv = query_embs['embeddings'][0].numpy().astype("float32").reshape(1, -1)
            print("Query embedding loaded successfully from .pt file.")
        except FileNotFoundError:
            print(f"Error: {query_file} not found. Encoding query from scratch.")
            qv = self.model.encode([query], normalize_embeddings=True)
            
        D, I = self.idx.search(qv, topk)
        items = []
        for d, i in zip(D[0], I[0]):
            if i < 0:
                continue
            it = dict(self.chunks[i])
            it["score"] = float(d)
            it["source"] = "dense"
            items.append(it)
        return items

# 하이브리드 검색기 (Dense + BM25(vector))
def _minmax(arr):
    arr = np.array(arr, dtype=float)
    if arr.size == 0:
        return arr
    lo, hi = float(arr.min()), float(arr.max())
    if hi - lo < 1e-12:
        return np.zeros_like(arr)
    return (arr - lo) / (hi - lo)

class HybridRetriever:
    def __init__(self, dense_indexer, tokenizer=lambda s: s.split(), fusion='rrf', alpha=0.6):
        self.dense = dense_indexer
        self.chunks = self.dense.chunks
        self.corpus = [c["text"] for c in self.chunks]
        self.tok = tokenizer
        self.fusion = fusion
        self.alpha = alpha
        self.bm25 = BM25Okapi([self.tok(t) for t in self.corpus]) if self.corpus else None

    def _rrf(self, dense_hits, bm_ranked, k=60, topk=10):
        scores = defaultdict(float)
        by_id = {}
        for lst in [dense_hits, bm_ranked]:
            for it in lst:
                if 'id' not in it:
                    it['id'] = id(it)
        
        for lst in [dense_hits, bm_ranked]:
            for r, it in enumerate(lst, start=1):
                _id = it.get("id")
                by_id[_id] = it
                scores[_id] += 1.0 / (k + r)
        fused = sorted(scores.items(), key=lambda x: x[1], reverse=True)[:topk]
        out = []
        for _id, s in fused:
            item = dict(by_id[_id])
            item["fused_score"] = float(s)
            item["source"] = "hybrid"
            out.append(item)
        return out

    def _weighted(self, dense_hits, bm_idx, bm_scores, topk=10):
        for it in dense_hits:
            if 'id' not in it:
                it['id'] = id(it)
        d_scores = [it.get("score", 0.0) for it in dense_hits]
        d_norm = _minmax(d_scores)
        d_map = {}
        for it, ns in zip(dense_hits, d_norm):
            d_map[it["id"]] = {"item": it, "norm": float(ns)}

        b_norm = _minmax(bm_scores)
        b_map = {}
        for i, ns in zip(bm_idx, b_norm):
            it = dict(self.chunks[i])
            it['id'] = id(it)
            it["source"] = "bm25"
            b_map[it["id"]] = {"item": it, "norm": float(ns)}

        merged_ids = set(d_map.keys()) | set(b_map.keys())
        heap = []
        for _id in merged_ids:
            dn = d_map.get(_id, {}).get("norm", 0.0)
            bn = b_map.get(_id, {}).get("norm", 0.0)
            fused = self.alpha * dn + (1.0 - self.alpha) * bn
            it = dict((d_map.get(_id) or b_map.get(_id))["item"])
            it["fused_score"] = float(fused)
            it["source"] = "hybrid"
            heap.append(it)
        heap.sort(key=lambda x: x["fused_score"], reverse=True)
        return heap[:topk]

    def search(self, query, k_dense=50, k_bm25=50, topk=10, rrf_k=60):
        dense_hits = self.dense.search_dense(query, topk=k_dense) if hasattr(self.dense, "search_dense") else []

        bm_ranked = []
        bm_idx = []
        bm_scores = []
        if self.bm25 is not None and self.corpus:
            qtok = self.tok(query)
            scores = self.bm25.get_scores(qtok)
            order = np.argsort(scores)[::-1][:k_bm25]
            bm_idx = order.tolist()
            bm_scores = scores[order].tolist()
            for i in bm_idx:
                it = dict(self.chunks[i])
                it["source"] = "bm25"
                bm_ranked.append(it)

        if self.fusion == 'rrf':
            return self._rrf(dense_hits, bm_ranked, k=rrf_k, topk=topk)
        else:
            return self._weighted(dense_hits, bm_idx, bm_scores, topk=topk)

#  모델 로드
model_name = "skt/A.X-4.0-Light"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForCausalLM.from_pretrained(
    model_name,
    torch_dtype="auto"
).to("cuda:1")
generator = pipeline("text-generation", model=model, tokenizer=tokenizer)

# RAG 인덱스 준비 (JSON 파일 로드)
JSON_FILES = ["rag_chunks.json"]
combined_text = ""
for json_file in JSON_FILES:
    try:
        with open(json_file, 'r', encoding='utf-8') as f:
            parsed_data = json.load(f)
            if isinstance(parsed_data, dict) and 'text' in parsed_data:
                combined_text += parsed_data['text'] + "\n\n"
            else:
                combined_text += str(parsed_data) + "\n\n"
            print(f"'{json_file}' 파일이 성공적으로 로드되었습니다.")
    except FileNotFoundError:
        print(f"오류: '{json_file}' 파일이 존재하지 않습니다. 파일을 확인해주세요.")
        exit()
    except json.JSONDecodeError:
        print(f"오류: '{json_file}' 파일의 JSON 형식이 올바르지 않습니다.")
        exit()

# ===== 섹션 앵커 분할 =====
ANCHOR_PATTERNS = [
    r"^\s*[IVXLC]+\.\s.+$",
    r"^\s*\d+\.\s.+$",
    r"^\s*\d+\-\d+\.\s.+$",
    r"^\s*제\s*\d+\s*기.*$",
    r"^\s*\(\d+\)\s.+$",
]
def detect_anchors(text: str):
    lines = text.splitlines()
    anchors = []
    for i, line in enumerate(lines):
        for pat in ANCHOR_PATTERNS:
            if re.match(pat, line.strip()):
                anchors.append((i, line.strip()))
                break
    return anchors

def segment_by_anchors(text: str):
    lines = text.splitlines()
    anchors = detect_anchors(text)
    if not anchors:
        return [{"title": "FULL", "text": text}]
    segments = []
    for idx, (lineno, title) in enumerate(anchors):
        start = lineno
        end = anchors[idx+1][0] if idx+1 < len(anchors) else len(lines)
        seg_text = "\n".join(lines[start:end]).strip()
        segments.append({"title": title, "text": seg_text})
    return segments

segments = segment_by_anchors(combined_text)
indexer = TextIndexer()
chunks = indexer.chunk(segments, max_chars=1000)
# cde_minicorpus.pt를 사용하여 빌드
indexer.build(embeddings_file="cde_minicorpus.pt")

retriever = HybridRetriever(
    indexer,
    tokenizer=lambda s: s.split(),
    fusion='rrf',
    alpha=0.6
)

# 가이드라인 JSON 파일 로드
GUIDELINE_FILE = "rnd_guideline.json"
guidelines = {}
try:
    with open(GUIDELINE_FILE, 'r', encoding='utf-8') as f:
        guidelines = json.load(f)
    print(f"'{GUIDELINE_FILE}' 파일이 성공적으로 로드되었습니다.")
except FileNotFoundError:
    print(f"오류: '{GUIDELINE_FILE}' 파일이 존재하지 않습니다. 파일을 확인해주세요.")
    exit()
except json.JSONDecodeError:
    print(f"오류: '{GUIDELINE_FILE}' 파일의 JSON 형식이 올바르지 않습니다.")
    exit()
# 사용자 입력 받기(R&D 과제명에서 핵심 키워드 추출)
depart_name = input("세부사업명: ")
project_no = input("연구개발 과제번호: ")
project_name = input("연구개발과제명: ")
period = input("전체 연구개발기간: ")
budget = input("총 연구비: (정부지원연구개발비 : 00,000천원, 기관부담 연구개발비 :00,000천원, 지방자치단체지원연구개발비 : 00,000천원): ")

#  섹션별 롤플레이 프롬프트 정의
section_roles = {
    "연구개발 목표": "당신은 R&D PMO입니다. 단계/일괄 협약의 최종목표를 500자 내외로 명확·간결·정량화하여 작성합니다. 핵심 성능지표(KPI), 달성 기준(수치/단위/마일스톤), 검증 방법을 포함하고 모호한 표현은 배제합니다.",
    "연구개발 내용": "당신은 기술 총괄(Tech Lead)입니다. 전체 연구범위를 1,000자 내외로 구조화해 기술 요소, 서브태스크, 인터페이스, 데이터/시스템 흐름을 설명하고 표준·규격·평가계획을 명시합니다.",
    "연구개발성과 활용계획 및 기대효과": "당신은 사업전략/사업개발(BD) 담당자입니다. 수요처, 적용 시나리오, 도입·확산 경로, 수익/비용 구조, 경제적 파급효과를 500자 내외로 정량·정성 지표와 함께 제시합니다.",
    "연구기획과제의 개요": "당신은 제안서 총괄 에디터입니다. 목적·필요성·기대효과를 일관된 논리로 요약해 과제가 해결하는 문제와 중요성을 한눈에 보이게 작성합니다.",
    "연구개발과제의 배경": "당신은 정책/RFP 적합성 분석가입니다. 관련 선행연구·시장/기술 동향·정부 정책·RFP/품목요약서 부합성을 근거와 함께 정리하고 제안 맥락을 명확히 합니다.",
    "연구개발과제의 필요성": "당신은 산업분석가입니다. 현황·문제점·시장규모/성장률·규제 및 정책 요구를 데이터로 제시하고, 해결 필요성을 인과적으로 설득력 있게 제시합니다.",
    "보안등급의 분류 및 해당 사유": "당신은 보안관리 책임자입니다. 국가연구개발혁신법 시행령 제45조 및 산업기술혁신사업 보안관리요령 제9조 기준을 근거로 보안등급과 결정 사유를 간결히 기재합니다.",
    "기술개발 핵심어(키워드)": "당신은 표준/용어 관리자입니다. 과제의 핵심 용어 5개를 한글/영문 정식 명칭으로 제시하고, 표준(협회/학회) 정의에 부합하도록 작성합니다.",
    "연차별 개발목표": "당신은 일정/성과관리 PM입니다. 연차별(1년차~n년차) 목표를 기관(주관/공동/참여 연구원)별로 구분해 KPI·마일스톤·검증기준을 정량화하여 제시합니다.",
    "연차별 개발내용 및 범위": "당신은 공동연구 컨소시엄 코디네이터입니다. 기관별 역할·범위·인계·의존성을 명확히 기술하고 중복/누락 없이 연차별 산출물과 책임을 표로 정리합니다(공동기관 없으면 생략).",
    "추진방법 및 전략": "당신은 기술전략/실험 설계 책임자입니다. 방법론(데이터·알고리즘·장비), 리스크와 대응책, 실험/검증 계획(평가지표·샘플수·통계/검증 절차)을 구체적으로 기술합니다.",
    "과제 성과의 활용방안": "당신은 제품/사업화 매니저입니다. 성과의 적용 분야, 기술 파급효과, 에너지 절감·환경 개선 등 기술적·사회적 효익을 사용 시나리오와 함께 제시합니다.",
    "신규사업 신설의 기대효과": "당신은 전략기획 임원입니다. 시장 창출, 일자리, 수입대체, 수출 증대, 비용 절감 등 경제·산업적 효과를 정량 지표(금액, 비율, 기간)와 함께 제시합니다.",
    "사회적 가치 창출 계획": "당신은 ESG/사회가치 책임자입니다. 개요-비전-목표-세부계획-기대효과 체계로 13개 사회적 가치 범주와의 연계를 명확히 하고 측정 가능한 지표를 포함합니다.",
    "사회적 가치창출의 기대효과": "당신은 임팩트 평가자입니다. 보건·안전·포용·지역·환경·민주성 등 사회적 가치 지표를 중심으로 성과/파급효과를 정량·정성으로 제시합니다.",
    "경제적 성과창출의 기대효과": "당신은 재무 담당자입니다(기업 작성). 매출/원가/영업이익, ROI/NPV, 고용효과 등 재무적 성과 전망을 가정과 산식(간단) 포함하여 명료하게 제시합니다.",
    "신규 인력 채용 계획 및 활용 방안": "당신은 HR 책임자입니다. 신규/기존 채용 구분, 채용 시점·역할·배치·활용 계획, 역량 매핑과 교육/온보딩 계획을 일정표와 함께 제시합니다."
}

#  섹션별 RAG 질의 템플릿 (JSON 파일에서 키워드 추출)
keywords = ""
if project_name:
    for item in guidelines.get("keywords", []):
        if item["project_name"] == project_name:
            keywords = item.get("keywords", "")
            break
section_queries = {
    "연구개발 목표": "최종목표(단계/일괄 협약목표)를 과제의 연구기획목표를 500자 내외로 기재합니다.",
    "연구개발 내용": "전체내용을 1,000자 내외로 기재합니다.",
    "연구개발성과 활용계획 및 기대효과": "연구기획의 수요처, 활용내용, 경제적 파급효과 등을 500자 내외로 기재합니다(연구시설ㆍ장비 구축을 목적으로 하는 과제의 경우에 연구시설ㆍ장비를 활용한 성과관리 및 자립운영계획, 수입금 관리 및 운영계획 등).",
    "연구기획과제의 개요": "연구기획과제의 개요는 연구개발과제의 목적, 필요성, 기대효과 등을 종합적으로 고려하여 작성합니다. 이를 통해 연구개발과제가 어떤 문제를 해결하고자 하는지, 왜 중요한지 명확히 제시합니다.",
    "연구개발과제의 배경": "구개발과제와 관련되는 연구개발과제의 배경 및 필요성, 정부 정책 및 RFP/품목요약서의 부합성 등을 종합적으로 기재합니다. 이를 통해 연구개발과제가 어떤 맥락에서 제안되었는지, 어떤 문제를 해결하고자 하는지 명확히 제시합니다.",
    "연구개발과제의 필요성": "연구개발과제의 필요성은 해당 기술 및 산업의 현황, 문제점, 시장 동향, 정책적 요구사항 등을 종합적으로 고려하여 작성합니다. 이를 통해 연구개발과제가 왜 필요한지, 어떤 문제를 해결하고자 하는지 명확히 제시합니다.",
    "보안등급의 분류 및 해당 사유": "국가연구개발혁신법 시행령 제45조(연구개발과제에 대한 보안과제의 분류) 및 산업기술혁신사업 보안관리요령 제9조(보안등급 분류 기준)을 참조하여, 계획서 표지에 있는 보안등급 분류에 대한 결정사유 기입합니다.",
    "기술개발 핵심어(키워드)": "핵심어는 동일 개발과제의 핵심적 용어로써 과제 관련 특수 용어로써 관련 업계, 협회나 학회 등에서 표준화되어 정의되었거나 일반화된 정식 명칭을 기재하며, 5개 단어를 한글 및 영문으로 반드시 기입하여야 함",
    "연차별 개발목표" : "1년차도, 2년차도, n년차도 각각의 연차별 개발목표를 주관연구개발기관, 공동연구개발기관, 참여연구원별로 구분하여 작성합니다.",
    "연차별 개발내용 및 범위" : "주관연구개발기관 및 공동연구개발기관이 담당하는 부분을 기술․표시하고, 연구개발기관별 연차별 개발목표, 내용 및 범위가 명확히 드러나도록 기술(공동연구개발기관이 없는 경우 생략)합니다.",
    "추진방법 및 전략" : "개발목표 달성을 위하여 무엇을 활용하고 어떻게 수행할 것인지 등 수행 방법을 구체적으로 기술하고, 세부개발 내용별 수행 방법, 수행 과정 중 예측되는 장애 요소 및 그 해결 방안, 계획된 실험과정 등을 기술합니다.",
    "과제 성과의 활용방안" : "연구개발과제 수행에 따라 예상되는 성과와 그 활용분야 및 활용방안을 기재하고, 기술적 측면은 해당 기술의 향상, 다른 기술로의 파급 효과 및 기술개발에 따른 에너지 절약 또는 환경 개선 효과 등을 서술합니다.",
    "신규사업 신설의 기대효과" : "연구개발성과의 과학ㆍ기술적, 경제ㆍ산업적, 사회적 측면에서 기대효과ㆍ파급효과 등을 기재하고, 경제․산업적 측면에는 시장 창출 및 일자리 창출 효과, 수입 대체 효과, 수출 증대 효과, 비용 절감 등의 경제적 효과와 산업발전에의 영향 등 산업적 효과를 서술합니다.",
    "사회적 가치 창출 계획" : "개요, 추진전략(비전), 추진목표, 세부계획, 기대효과로 구분하여 작성합니다.사회적 가치란 사회적·경제적·환경적·문화적 영역에서 공공의 이익과 공동체 발전에 기여하는 가치로서 13가지(인간의 존엄성을 유지하는 기본권리로서 인권보호, 재난과 사고로부터 안전한 근로 생활환경의 유지, 건강한 생활이 가능한 보건복지의 제공, 노동권의 보장과 근로조건의 향상, 사회적 약자에 대한 기회제공과 사회통합, 대기업·중소기업간 상생과 협력, 품위있는 삶을 누릴 수 있는 양질의 일자리 창출, 지역사회 활성화와 공동체 복원, 경제활동을 통한 이익이 지역에 순환되는 지역경제 공헌, 윤리적 생산과 유통을 포함한 기업의 자발적인 사회적 책임 이행, 환경의 지속가능성 보전, 시민적 권리로서 민주적 의사결정과 참여의 실현, 그 밖에 공동체의 이익 실현과 공공성 강화)을 포괄하는 가치를 의미합니다.",
    "사회적 가치창출의 기대효과": "연구개발과제 수행에 따라 예상되는 성과와 그 활용분야 및 활용방안을 기재하고, 기술적 측면은 해당 기술의 향상, 다른 기술로의 파급 효과 및 기술개발에 따른 에너지 절약 또는 환경 개선 효과 등을 서술합니다.",
    "경제적 성과창출의 기대효과" : "주관/공동연구개발기관 중 기업만 작성합니다.",
    "신규 인력 채용 계획 및 활용 방안" : "신규 채용 여부는 신규 채용인 경우와 기존인 경우로 표기합니다. 또한, 신규 채용 구분 여부는 동 과제 수행을 위해 사업 공고일 기준 6개월 이전에 신규로 채용했거나 과제 전체 연구개발기간 중 채용 계획이 있는 경우로 구분하여 서술합니다.",
}

# 자동 문장 생성 함수 (RAG 근거 포함으로 변경)
def search_contexts(section: str, topk=4):
    query = section_queries.get(section, section)
    hits = retriever.search(query, k_dense=4, k_bm25=4)
    contexts = []
    max_ctx_len = 900
    for h in hits:
        t = h["text"]
        if len(t) > max_ctx_len:
            t = t[:max_ctx_len] + "..."
        contexts.append(t)
    return contexts

def build_prompt_with_context(section, role_instruction, contexts):
    ctx_block = "\n\n".join([f"[근거]\n{c}" for c in contexts]) if contexts else "[근거]\n(해당 섹션에 대한 근거 스니펫 없음)"
    prompt = f"""
역할: {role_instruction}
작성 항목: [{section}]
세부사업명: {depart_name}
연구개발 과제번호: {project_no}

작성 조건:
    - 제시된 {GUIDELINE_FILE} 가이드라인을 엄격히 준수하여 작성합니다.
    - R&D 결과물과 기술적으로 직접적인 연관성이 적은 용어나 화려한 미사여구(고부가가치, 차세대, 첨단 등)는 사용을 삼가야 하여 작성합니다.
    - 구체적인 규격이나 범위를 함께 활용하여 작성합니다.
    - {ctx_block}
    - {section} 작성 시 위 근거를 반드시 반영합니다.
    - 반드시 {JSON_FILES} 작성방식과 구성을 참고하여 작성합니다.
    - 문단마다 핵심 키워드를 포함하여 작성합니다.
    - 문장 길이는 다양하게 구성합니다.
    - 문장 시작은 다양하게 합니다.
    - 다른 항목과 중복되는 표현은 피합니다.
    - 각 항목의 특성에 맞게 문장 및 문단을 작성합니다.
    - 전문적이면서 친화적인 톤을 사용하여 작성합니다.
    - 연구개발계획서를 쉽게 이해할 수 있도록 모든 전문용어 또는 약어에 대한 주석처리(약어는 full name 표기)를 합니다.
    """

    return prompt.strip()

def generate_text(section, keywords=""):
    role_instruction = section_roles.get(section, "")
    contexts = search_contexts(section, topk=4)
    prompt = build_prompt_with_context(section, role_instruction, contexts)
    output = generator(
        prompt,
        max_new_tokens=5000,
        do_sample=False,
        temperature=0,
        top_p=0.9,
        repetition_penalty=1.05,
        eos_token_id=tokenizer.eos_token_id
    )
    text = output[0]["generated_text"]
    gen = text[len(prompt):].strip() if text.startswith(prompt) else text.strip()
    return gen

#  DOCX 문서 생성 및 서식 지정
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
font.size = Pt(11)

doc.add_heading("연구개발계획서", 0)
doc.add_paragraph(f"세부사업명: {depart_name}")
doc.add_paragraph(f"연구개발 과제번호: {project_no}")
doc.add_paragraph(f"연구개발과제명: {project_name}")
doc.add_paragraph(f"전체 연구개발기간: {period}")
doc.add_paragraph(f"예산: {budget} 천원")
doc.add_paragraph("")

sections = ["연구개발 목표", "연구개발 내용", "연구개발성과 활용계획 및 기대효과",
            "연구기획과제의 개요", "연구개발과제의 배경", "연구개발과제의 필요성",
            "보안등급의 분류 및 해당 사유", "기술개발 핵심어(키워드)", "연차별 개발목표",
            "연차별 개발내용 및 범위", "추진방법 및 전략", "과제 성과의 활용방안",
            "신규사업 신설의 기대효과", "사회적 가치 창출 계획", "사회적 가치창출의 기대효과",
            "경제적 성과창출의 기대효과", "신규 인력 채용 계획 및 활용 방안"]

for section in sections:
    doc.add_heading(section, level=1)
    doc.add_paragraph(generate_text(section, keywords))
doc.add_page_break()

# 파일 저장
output_file = "연구개발계획서(test_ver).docx"
doc.save(output_file)
print(f"완료: '{output_file}' 파일이 생성되었습니다!")